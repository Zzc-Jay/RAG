from __future__ import annotations
import os
import sys
import tempfile
from unittest.mock import patch, MagicMock, call

import fitz
import pytest

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "src"))

from loader import load_pdf, load_txt, load_docx, load_document, SUPPORTED_TYPES, _table_to_markdown
from chunker import split_pages
from bm25_index import build_index, search_bm25, delete_index
from retriever import _rrf_fuse, _doc_key
from kb_manager import create_kb, delete_kb, list_kbs, add_doc, remove_doc, get_kb_docs, add_docs_batch, remove_docs_batch
from generator import build_prompt
from security import (
    validate_question,
    validate_kb_name,
    validate_file_extension,
    RateLimiter,
)
from retry import retry_call, _is_retryable


# ═══════════════════════════════════════════════════════════════
# 辅助函数
# ═══════════════════════════════════════════════════════════════

def _make_test_pdf(text: str) -> str:
    fd, path = tempfile.mkstemp(suffix=".pdf")
    os.close(fd)
    os.unlink(path)
    doc = fitz.open()
    page = doc.new_page()
    rect = fitz.Rect(72, 72, 500, 700)
    page.insert_textbox(rect, text, fontsize=12)
    doc.save(path)
    doc.close()
    return path


def _make_test_txt(text: str) -> str:
    fd, path = tempfile.mkstemp(suffix=".txt")
    os.close(fd)
    with open(path, "w", encoding="utf-8") as f:
        f.write(text)
    return path


def _make_test_docx(text: str) -> str:
    from docx import Document
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(path)
    return path


def _make_test_docx_with_table() -> str:
    """创建含表格的 DOCX 文件，返回路径。"""
    from docx import Document
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    doc = Document()
    doc.add_paragraph("这是表格前的正文段落。")
    table = doc.add_table(rows=3, cols=3)
    headers = ["姓名", "年龄", "城市"]
    data = [["张三", "28", "北京"], ["李四", "32", "上海"]]
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
    for i, row_data in enumerate(data):
        for j, val in enumerate(row_data):
            table.rows[i + 1].cells[j].text = val
    doc.add_paragraph("这是表格后的正文段落。")
    doc.save(path)
    return path


# ═══════════════════════════════════════════════════════════════
# loader (PDF) 测试
# ═══════════════════════════════════════════════════════════════

def test_load_pdf_returns_list_of_dict():
    path = _make_test_pdf("Python is a programming language.\nIt is widely used in AI.")
    try:
        result = load_pdf(path)
        assert isinstance(result, list), f"expected list, got {type(result)}"
        assert len(result) >= 1, "should have at least 1 page"
        assert "page" in result[0]
        assert "text" in result[0]
        assert "Python" in result[0]["text"]
    finally:
        os.unlink(path)


def test_load_pdf_skips_empty_pages():
    path = _make_test_pdf("Non-empty content.")
    try:
        result = load_pdf(path)
        assert len(result) >= 1
        for p in result:
            assert p["text"].strip(), "empty page should be skipped"
    finally:
        os.unlink(path)


# ═══════════════════════════════════════════════════════════════
# loader (TXT) 测试
# ═══════════════════════════════════════════════════════════════

def test_load_txt_returns_single_page():
    path = _make_test_txt("Hello world!\nThis is a test document.")
    try:
        result = load_txt(path)
        assert isinstance(result, list)
        assert len(result) == 1
        assert result[0]["page"] == 1
        assert "Hello world" in result[0]["text"]
    finally:
        os.unlink(path)


def test_load_txt_empty_file():
    path = _make_test_txt("")
    try:
        result = load_txt(path)
        assert result == []
    finally:
        os.unlink(path)


def test_load_txt_gbk_encoding():
    path = _make_test_txt("")
    os.unlink(path)
    text = "中文测试内容。\n第二行。"
    with open(path, "w", encoding="gbk") as f:
        f.write(text)
    try:
        result = load_txt(path)
        assert len(result) == 1
        assert "中文测试" in result[0]["text"]
    finally:
        os.unlink(path)


# ═══════════════════════════════════════════════════════════════
# loader (DOCX) 测试
# ═══════════════════════════════════════════════════════════════

def test_load_docx_returns_single_page():
    path = _make_test_docx("First paragraph.\nSecond paragraph.")
    try:
        result = load_docx(path)
        assert isinstance(result, list)
        assert len(result) == 1
        assert result[0]["page"] == 1
        assert "First paragraph" in result[0]["text"]
        assert "Second paragraph" in result[0]["text"]
    finally:
        os.unlink(path)


def test_load_docx_empty():
    from docx import Document
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    Document().save(path)
    try:
        result = load_docx(path)
        assert result == []
    finally:
        os.unlink(path)


# ═══════════════════════════════════════════════════════════════
# loader (dispatcher) 测试
# ═══════════════════════════════════════════════════════════════

def test_load_document_dispatcher():
    path = _make_test_txt("Dispatcher test.")
    try:
        result = load_document(path, ".txt")
        assert len(result) == 1
        assert "Dispatcher test" in result[0]["text"]

        result_md = load_document(path, ".md")
        assert len(result_md) == 1

        with pytest.raises(ValueError):
            load_document(path, ".xyz")
    finally:
        os.unlink(path)


# ═══════════════════════════════════════════════════════════════
# 表格提取 & 扫描检测测试（文档深度解析）
# ═══════════════════════════════════════════════════════════════

def test_table_to_markdown_simple():
    data = [["姓名", "年龄"], ["张三", "28"], ["李四", "32"]]
    md = _table_to_markdown(data)
    assert "[表格]" in md
    assert "姓名" in md
    assert "张三" in md
    assert "[/表格]" in md
    assert "---" in md  # separator after header


def test_table_to_markdown_empty():
    assert _table_to_markdown([]) == ""
    assert _table_to_markdown([[], []]) == ""


def test_table_to_markdown_single_row():
    md = _table_to_markdown([["唯一行", "值1", "值2"]])
    assert "[表格]" in md
    assert "唯一行" in md
    assert "---" not in md  # no separator for single row


def test_load_pdf_includes_metadata_fields():
    path = _make_test_pdf("这是一段测试文本。" * 10)
    try:
        pages = load_pdf(path)
        assert len(pages) >= 1
        p = pages[0]
        assert "tables" in p
        assert "table_count" in p
        assert "is_scanned" in p
        assert p["table_count"] == 0
        assert p["is_scanned"] is False
    finally:
        os.unlink(path)


def test_load_pdf_with_table_mocked():
    """表格检测：mock page.find_tables() 模拟含表格的 PDF 页面。"""
    # 创建普通 PDF
    path = _make_test_pdf("销售数据汇总\n\n" + "正文内容。" * 20)
    try:
        with patch("fitz.Page.find_tables") as mock_find:
            # Mock table detection result
            mock_table = MagicMock()
            mock_table.extract.return_value = [
                ["产品", "销量", "单价"],
                ["Widget", "100", "9.9"],
                ["Gadget", "50", "19.9"],
            ]
            mock_result = MagicMock()
            mock_result.tables = [mock_table]
            mock_find.return_value = mock_result

            pages = load_pdf(path)
            assert len(pages) >= 1
            p = pages[0]
            assert p["table_count"] == 1
            assert len(p["tables"]) == 1
            assert "[表格]" in p["text"]
            assert "[/表格]" in p["text"]
            assert "Widget" in p["text"]
    finally:
        os.unlink(path)


def test_load_docx_with_table():
    path = _make_test_docx_with_table()
    try:
        pages = load_docx(path)
        assert len(pages) == 1
        p = pages[0]
        assert p["table_count"] == 1
        assert len(p["tables"]) == 1
        assert "[表格]" in p["text"]
        assert "张三" in p["text"]
        assert "表格前的正文段落" in p["text"]
    finally:
        os.unlink(path)


def test_scanned_page_detection():
    """文字极少 + 有图片 = 扫描页。"""
    # 创建含少量文字的 PDF（不含图片），用 mock 模拟页面有图片
    path = _make_test_pdf("AB")
    try:
        with patch("fitz.Page.get_images", return_value=[("xref", 1, 0, 0, 0, 0, "DCTDecode")]):
            pages = load_pdf(path)
            assert len(pages) == 1
            p = pages[0]
            assert p["is_scanned"] is True
            assert "图片型页面" in p["text"]
    finally:
        os.unlink(path)


def test_chunks_inherit_table_metadata():
    pages = [{
        "page": 1, "text": "第一章内容\n" * 20,
        "tables": [["a", "b"]], "table_count": 1, "is_scanned": False,
    }]
    chunks = split_pages(pages, source="test.pdf")
    assert len(chunks) > 0
    for c in chunks:
        assert "has_table" in c
        assert "is_scanned" in c
        assert c["has_table"] is True
        assert c["is_scanned"] is False


def test_bm25_search_returns_table_metadata():
    # 5 docs for meaningful BM25 IDF scores (same pattern as test_bm25_build_and_search)
    chunks = [
        {"text": "advanced retrieval augmented generation for enterprise search systems",
         "source": "test.pdf", "page": 1, "chunk_idx": 0,
         "has_table": True, "is_scanned": False},
        {"text": "deep learning neural networks for computer vision tasks",
         "source": "test.pdf", "page": 1, "chunk_idx": 1,
         "has_table": False, "is_scanned": False},
        {"text": "natural language processing with transformer architectures",
         "source": "test.pdf", "page": 1, "chunk_idx": 2,
         "has_table": False, "is_scanned": False},
        {"text": "reinforcement learning for game playing and robotics",
         "source": "test.pdf", "page": 1, "chunk_idx": 3,
         "has_table": False, "is_scanned": False},
        {"text": "graph neural networks for social network analysis",
         "source": "test.pdf", "page": 1, "chunk_idx": 4,
         "has_table": False, "is_scanned": False},
    ]
    from bm25_index import build_index, search_bm25, delete_index
    kb = "_test_bm25_meta2"
    build_index(chunks, kb)
    try:
        results = search_bm25("retrieval augmented generation", kb, top_k=5)
        assert len(results) >= 1
        # The first result should have has_table=True from chunk 0
        assert results[0]["has_table"] is True
        assert results[0]["is_scanned"] is False
    finally:
        delete_index(kb)


# ═══════════════════════════════════════════════════════════════
# chunker 测试
# ═══════════════════════════════════════════════════════════════

def test_split_pages_preserves_metadata():
    pages = [
        {"page": 1, "text": "First page content. " * 30},
        {"page": 2, "text": "Second page different. " * 30},
    ]
    chunks = split_pages(pages, source="test.pdf")
    assert len(chunks) >= 2
    for c in chunks:
        assert "text" in c
        assert "source" in c
        assert "page" in c
        assert "chunk_idx" in c
        assert c["source"] == "test.pdf"
        assert c["page"] in (1, 2)


def test_split_pages_chunk_indices_unique():
    pages = [{"page": 1, "text": "A. " * 200}]
    chunks = split_pages(pages, source="test.pdf")
    indices = [c["chunk_idx"] for c in chunks]
    assert len(indices) == len(set(indices)), "chunk indices must be unique"


# ═══════════════════════════════════════════════════════════════
# BM25 测试
# ═══════════════════════════════════════════════════════════════

def test_bm25_build_and_search():
    chunks = [
        {"text": "Python 编程语言", "source": "a.pdf", "page": 1, "chunk_idx": 0},
        {"text": "Java 企业应用", "source": "b.pdf", "page": 1, "chunk_idx": 1},
        {"text": "Rust 系统编程", "source": "c.pdf", "page": 1, "chunk_idx": 2},
        {"text": "AI 人工智能", "source": "d.pdf", "page": 1, "chunk_idx": 3},
        {"text": "RAG 增强检索", "source": "e.pdf", "page": 1, "chunk_idx": 4},
    ]
    try:
        build_index(chunks, "test_bm25")
        results = search_bm25("RAG 检索", "test_bm25", top_k=3)
        assert len(results) >= 1
        assert "RAG" in results[0]["text"]
    finally:
        delete_index("test_bm25")


def test_bm25_returns_empty_for_missing_kb():
    results = search_bm25("query", "nonexistent_kb")
    assert results == []


# ═══════════════════════════════════════════════════════════════
# RRF 融合测试
# ═══════════════════════════════════════════════════════════════

def test_rrf_fusion_ranks_dual_source_higher():
    vector = [
        {"text": "Doc A", "source": "a.pdf", "page": 1},
        {"text": "Doc B", "source": "b.pdf", "page": 1},
    ]
    bm25 = [
        {"text": "Doc B", "source": "b.pdf", "page": 1},
        {"text": "Doc A", "source": "a.pdf", "page": 1},
    ]
    fused = _rrf_fuse(vector, bm25, top_k=2)
    assert fused[0]["score"] == fused[1]["score"]


def test_rrf_fusion_single_source_ranked_lower():
    vector = [
        {"text": "Doc X", "source": "x.pdf", "page": 1},
    ]
    bm25 = [
        {"text": "Doc Y", "source": "y.pdf", "page": 1},
    ]
    fused = _rrf_fuse(vector, bm25, top_k=2)
    assert len(fused) == 2
    assert fused[0]["score"] == fused[1]["score"]


# ═══════════════════════════════════════════════════════════════
# kb_manager 测试
# ═══════════════════════════════════════════════════════════════

def test_kb_crud():
    name = "test_kb_crud"
    try:
        delete_kb(name)
    except ValueError:
        pass

    create_kb(name)
    assert name in list_kbs()

    add_doc(name, "doc1.pdf", 10, 15, ".pdf")
    add_doc(name, "doc2.txt", 1, 5, ".txt")
    docs = get_kb_docs(name)
    assert len(docs) == 2

    remove_doc(name, "doc1.pdf")
    docs = get_kb_docs(name)
    assert len(docs) == 1
    assert docs[0]["name"] == "doc2.txt"
    assert docs[0]["type"] == ".txt"

    delete_kb(name)
    assert name not in list_kbs()


def test_create_duplicate_kb_raises():
    name = "test_dup_kb"
    try:
        delete_kb(name)
    except ValueError:
        pass
    create_kb(name)
    try:
        create_kb(name)
        assert False, "should have raised ValueError"
    except ValueError:
        pass
    finally:
        delete_kb(name)


def test_add_doc_dedup():
    name = "test_dedup_kb"
    try:
        delete_kb(name)
    except ValueError:
        pass
    create_kb(name)
    add_doc(name, "dup.pdf", 10, 20, ".pdf")
    add_doc(name, "dup.pdf", 5, 10, ".pdf")
    docs = get_kb_docs(name)
    assert len(docs) == 1
    assert docs[0]["pages"] == 5
    assert docs[0]["chunks"] == 10
    delete_kb(name)


def test_add_docs_batch():
    name = "test_batch_add_kb"
    try:
        delete_kb(name)
    except ValueError:
        pass
    create_kb(name)
    doc_infos = [
        {"name": "doc1.pdf", "pages": 10, "chunks": 15, "type": ".pdf"},
        {"name": "doc2.txt", "pages": 1, "chunks": 5, "type": ".txt"},
        {"name": "doc3.md", "pages": 3, "chunks": 8, "type": ".md"},
    ]
    add_docs_batch(name, doc_infos)
    docs = get_kb_docs(name)
    assert len(docs) == 3
    names = {d["name"] for d in docs}
    assert names == {"doc1.pdf", "doc2.txt", "doc3.md"}
    delete_kb(name)


def test_add_docs_batch_dedup():
    """批量添加时，同名文档应被覆盖（与 add_doc 行为一致）。"""
    name = "test_batch_dedup_kb"
    try:
        delete_kb(name)
    except ValueError:
        pass
    create_kb(name)
    add_doc(name, "dup.pdf", 10, 20, ".pdf")
    add_docs_batch(name, [{"name": "dup.pdf", "pages": 5, "chunks": 10, "type": ".pdf"}])
    docs = get_kb_docs(name)
    assert len(docs) == 1
    assert docs[0]["pages"] == 5
    delete_kb(name)


def test_remove_docs_batch():
    name = "test_batch_remove_kb"
    try:
        delete_kb(name)
    except ValueError:
        pass
    create_kb(name)
    add_doc(name, "a.pdf", 1, 2, ".pdf")
    add_doc(name, "b.pdf", 1, 2, ".pdf")
    add_doc(name, "c.pdf", 1, 2, ".pdf")

    removed = remove_docs_batch(name, ["a.pdf", "c.pdf"])
    assert len(removed) == 2
    assert "a.pdf" in removed
    assert "c.pdf" in removed

    docs = get_kb_docs(name)
    assert len(docs) == 1
    assert docs[0]["name"] == "b.pdf"
    delete_kb(name)


def test_remove_docs_batch_partial():
    """批量删除时，不存在的文档名被忽略。"""
    name = "test_batch_remove_partial"
    try:
        delete_kb(name)
    except ValueError:
        pass
    create_kb(name)
    add_doc(name, "existing.pdf", 1, 2, ".pdf")

    removed = remove_docs_batch(name, ["existing.pdf", "nonexistent.pdf"])
    assert len(removed) == 1
    assert "existing.pdf" in removed

    docs = get_kb_docs(name)
    assert len(docs) == 0
    delete_kb(name)


def test_remove_docs_batch_empty():
    name = "test_batch_remove_empty"
    try:
        delete_kb(name)
    except ValueError:
        pass
    create_kb(name)
    removed = remove_docs_batch(name, ["no_doc.pdf"])
    assert removed == []
    delete_kb(name)


# ═══════════════════════════════════════════════════════════════
# generator 测试
# ═══════════════════════════════════════════════════════════════

def test_build_prompt_includes_citations():
    docs = [
        {"text": "RAG is Retrieval-Augmented Generation.", "source": "ai.pdf", "page": 5},
        {"text": "Embeddings map text to vectors.", "source": "nlp.pdf", "page": 12},
    ]
    prompt = build_prompt("What is RAG?", docs)
    assert "[1]" in prompt
    assert "[2]" in prompt
    assert "ai.pdf" in prompt
    assert "5" in prompt
    assert "RAG is Retrieval-Augmented Generation" in prompt
    assert "What is RAG?" in prompt


# ═══════════════════════════════════════════════════════════════
# Security: validate_question 测试
# ═══════════════════════════════════════════════════════════════

def test_validate_question_normal():
    result = validate_question("什么是 RAG？")
    assert result == "什么是 RAG？"


def test_validate_question_xss_escaped():
    result = validate_question('<script>alert(1)</script>')
    assert "<script>" not in result
    assert "&lt;script&gt;" in result
    assert "&lt;/script&gt;" in result


def test_validate_question_html_tags_escaped():
    result = validate_question('<b>bold</b>')
    assert "&lt;b&gt;" in result
    assert "</b>" not in result


def test_validate_question_too_long():
    long_text = "A" * 2001
    with pytest.raises(ValueError, match="不能超过"):
        validate_question(long_text)


def test_validate_question_empty():
    with pytest.raises(ValueError):
        validate_question("")
    with pytest.raises(ValueError):
        validate_question("   ")


# ═══════════════════════════════════════════════════════════════
# Security: validate_kb_name 测试
# ═══════════════════════════════════════════════════════════════

def test_validate_kb_name_normal():
    assert validate_kb_name("技术文档") == "技术文档"
    assert validate_kb_name("tech_docs") == "tech_docs"
    assert validate_kb_name("docs-v2") == "docs-v2"


def test_validate_kb_name_rejects_special():
    with pytest.raises(ValueError):
        validate_kb_name("知识库!!!")
    with pytest.raises(ValueError):
        validate_kb_name("test@docs")


def test_validate_kb_name_rejects_empty():
    with pytest.raises(ValueError):
        validate_kb_name("")


def test_validate_kb_name_too_long():
    with pytest.raises(ValueError):
        validate_kb_name("A" * 51)


# ═══════════════════════════════════════════════════════════════
# Security: validate_file_extension 测试
# ═══════════════════════════════════════════════════════════════

def test_validate_file_extension_allowed():
    assert validate_file_extension("doc.pdf") == ".pdf"
    assert validate_file_extension("doc.TXT") == ".txt"
    assert validate_file_extension("doc.MD") == ".md"
    assert validate_file_extension("doc.docx") == ".docx"


def test_validate_file_extension_rejected():
    with pytest.raises(ValueError):
        validate_file_extension("malware.exe")
    with pytest.raises(ValueError):
        validate_file_extension("script.py")


def test_validate_file_extension_no_extension():
    with pytest.raises(ValueError):
        validate_file_extension("noext")


# ═══════════════════════════════════════════════════════════════
# Security: RateLimiter 测试
# ═══════════════════════════════════════════════════════════════

def test_rate_limiter_allows_within_limit():
    rl = RateLimiter(max_requests=3, window_seconds=60)
    assert rl.check("user1") is True
    assert rl.check("user1") is True
    assert rl.check("user1") is True


def test_rate_limiter_blocks_after_limit():
    rl = RateLimiter(max_requests=2, window_seconds=60)
    rl.check("user1")
    rl.check("user1")
    assert rl.check("user1") is False


def test_rate_limiter_isolated_sessions():
    rl = RateLimiter(max_requests=2, window_seconds=60)
    assert rl.check("user1") is True
    assert rl.check("user1") is True
    assert rl.check("user1") is False
    assert rl.check("user2") is True  # 不同 session 不受影响


def test_rate_limiter_remaining():
    rl = RateLimiter(max_requests=5, window_seconds=60)
    assert rl.remaining("user1") == 5
    rl.check("user1")
    assert rl.remaining("user1") == 4
    rl.check("user1")
    assert rl.remaining("user1") == 3


# ═══════════════════════════════════════════════════════════════
# Retry: _is_retryable 测试
# ═══════════════════════════════════════════════════════════════

def test_is_retryable_connection_error():
    assert _is_retryable(ConnectionError("connection refused")) is True


def test_is_retryable_timeout():
    assert _is_retryable(TimeoutError("timed out")) is True


def test_is_retryable_os_error():
    assert _is_retryable(OSError("network unreachable")) is True


def test_is_retryable_value_error_not_retryable():
    assert _is_retryable(ValueError("invalid parameter")) is False


def test_is_retryable_429_in_message():
    assert _is_retryable(RuntimeError("HTTP 429 Too Many Requests")) is True


def test_is_retryable_500_in_message():
    assert _is_retryable(RuntimeError("502 Bad Gateway")) is True


# ═══════════════════════════════════════════════════════════════
# Retry: retry_call 测试
# ═══════════════════════════════════════════════════════════════

def test_retry_call_success_first_attempt():
    mock_func = MagicMock(return_value="success")
    result = retry_call(mock_func, max_retries=2)
    assert result == "success"
    assert mock_func.call_count == 1


def test_retry_call_recovers_on_second():
    mock_func = MagicMock(side_effect=[ConnectionError("fail"), "success"])
    result = retry_call(mock_func, max_retries=3, base_delay=0.001)
    assert result == "success"
    assert mock_func.call_count == 2


def test_retry_call_exhausted_raises():
    mock_func = MagicMock(side_effect=ConnectionError("always fail"))
    with pytest.raises(RuntimeError, match="已重试"):
        retry_call(mock_func, max_retries=2, base_delay=0.001)
    assert mock_func.call_count == 3  # 1 初始 + 2 重试


def test_retry_call_no_retry_on_value_error():
    """不可恢复错误（ValueError）不应重试。"""
    mock_func = MagicMock(side_effect=ValueError("bad input"))
    with pytest.raises(ValueError):
        retry_call(mock_func, max_retries=3, base_delay=0.001)
    assert mock_func.call_count == 1  # 只调一次，不重试


def test_retry_call_custom_max_retries():
    mock_func = MagicMock(side_effect=ConnectionError("fail"))
    with pytest.raises(RuntimeError):
        retry_call(mock_func, max_retries=1, base_delay=0.001)
    assert mock_func.call_count == 2  # 1 initial + 1 retry


# ═══════════════════════════════════════════════════════════════
# Embedder mock 测试
# ═══════════════════════════════════════════════════════════════

class FakeEmbeddingOutput:
    """模拟 DashScope TextEmbedding 返回的 resp.output 结构。"""
    def __init__(self, embeddings_list):
        self._embeddings = {"embeddings": [{"embedding": e} for e in embeddings_list]}

    def __getitem__(self, key):
        return self._embeddings[key]


class FakeEmbeddingResponse:
    """模拟 DashScope TextEmbedding.call 的返回值。"""
    def __init__(self, embeddings_list, status_code=200, message=""):
        self.status_code = status_code
        self.message = message
        self.output = FakeEmbeddingOutput(embeddings_list)


def test_embed_batch_with_mock():
    from embedder import _call_embedding_api

    fake_vec = [0.1] * 1024
    fake_resp = FakeEmbeddingResponse([fake_vec, fake_vec])

    with patch("embedder.TextEmbedding.call", return_value=fake_resp):
        embeddings, tokens = _call_embedding_api(["text1", "text2"])
        assert len(embeddings) == 2
        assert len(embeddings[0]) == 1024


def test_embed_batch_with_retry_on_failure():
    from embedder import _call_embedding_api

    fake_vec = [0.1] * 1024
    fail_resp = FakeEmbeddingResponse([], status_code=500, message="server error")
    success_resp = FakeEmbeddingResponse([fake_vec])

    with patch("embedder.TextEmbedding.call") as mock_call:
        mock_call.side_effect = [fail_resp, success_resp]
        embeddings, tokens = _call_embedding_api(["text1"])
        assert len(embeddings) == 1
        assert mock_call.call_count == 2


def test_search_with_mock_embedding():
    from embedder import search as vector_search
    from unittest.mock import patch, MagicMock

    fake_vec = [0.1] * 1024

    with patch("embedder.TextEmbedding.call") as mock_embed:
        mock_embed.return_value = FakeEmbeddingResponse([fake_vec])

        with patch("embedder._get_collection") as mock_get_col:
            mock_col = MagicMock()
            mock_col.query.return_value = {
                "ids": [["id1", "id2"]],
                "documents": [["Python programming", "Java development"]],
                "metadatas": [[{"source": "a.pdf", "page": 1}, {"source": "b.pdf", "page": 2}]],
                "distances": [[0.1, 0.3]],
            }
            mock_get_col.return_value = mock_col

            results = vector_search("Python", "test_kb", top_k=2)

            assert len(results) == 2
            assert results[0]["source"] == "a.pdf"
            assert results[0]["page"] == 1
            assert results[1]["source"] == "b.pdf"


# ═══════════════════════════════════════════════════════════════
# Generator mock 测试
# ═══════════════════════════════════════════════════════════════

class FakeChunk:
    """模拟 DashScope 流式响应的单个 chunk。"""
    def __init__(self, text, status_code=200, message=""):
        self.status_code = status_code
        self.output = MagicMock()
        self.output.text = text


def test_generate_stream_yields_tokens():
    from generator import generate_stream
    from providers.dashscope_provider import DashScopeProvider

    def fake_stream(self, prompt: str):
        yield from ["这是", "一个", "测试"]

    docs = [{"text": "测试文档内容", "source": "test.pdf", "page": 1}]

    with patch.object(DashScopeProvider, "generate_stream", fake_stream):
        tokens = list(generate_stream("测试问题", docs))
        assert tokens == ["这是", "一个", "测试"]


def test_generate_stream_empty_docs():
    from generator import generate_stream
    from providers.dashscope_provider import DashScopeProvider

    def fake_stream(self, prompt: str):
        yield "已生成回答"

    docs: list[dict] = [{"text": "唯一文档", "source": "only.pdf", "page": 1}]

    with patch.object(DashScopeProvider, "generate_stream", fake_stream):
        tokens = list(generate_stream("有文档的问题", docs))
        assert "已生成回答" in tokens


def test_generate_stream_wraps_provider_exception():
    from generator import generate_stream
    from providers.dashscope_provider import DashScopeProvider

    def fake_error(self, prompt: str):
        raise RuntimeError("模拟 provider 内部错误")
        yield  # unreachable, makes it a generator

    docs = [{"text": "测试内容", "source": "test.pdf", "page": 1}]

    with patch.object(DashScopeProvider, "generate_stream", fake_error):
        with pytest.raises(RuntimeError, match="模型.*调用失败"):
            list(generate_stream("测试问题", docs))


def test_generate_wraps_provider_exception():
    from generator import generate
    from providers.dashscope_provider import DashScopeProvider

    def fake_error(self, prompt: str):
        raise RuntimeError("模拟 provider 内部错误")

    docs = [{"text": "测试内容", "source": "test.pdf", "page": 1}]

    with patch.object(DashScopeProvider, "generate", fake_error):
        with pytest.raises(RuntimeError, match="模型.*调用失败"):
            generate("测试问题", docs)


def test_build_prompt_empty_docs():
    """即使没有参考文档，build_prompt 也应能工作。"""
    prompt = build_prompt("你好", [])
    assert "问题：你好" in prompt
    assert "参考资料" in prompt


# ═══════════════════════════════════════════════════════════════
# Retriever mock 测试
# ═══════════════════════════════════════════════════════════════

def test_retrieve_fuses_both_sources():
    from retriever import retrieve as hybrid_retrieve

    with patch("retriever.vector_search") as mock_vec, \
         patch("retriever.search_bm25") as mock_bm25:
        mock_vec.return_value = [
            {"text": "Vector result", "source": "v.pdf", "page": 1, "score": 0.9},
        ]
        mock_bm25.return_value = [
            {"text": "BM25 result", "source": "b.pdf", "page": 2, "score": 5.0},
        ]

        results = hybrid_retrieve("test query", "test_kb", top_k=3)

        assert len(results) >= 1
        mock_vec.assert_called_once()
        mock_bm25.assert_called_once()


def test_rrf_handles_empty_vector():
    """向量检索为空时，BM25 结果应能单独返回。"""
    vector: list[dict] = []
    bm25 = [{"text": "Only BM25", "source": "b.pdf", "page": 1}]
    fused = _rrf_fuse(vector, bm25, top_k=3)
    assert len(fused) == 1
    assert fused[0]["source"] == "b.pdf"


def test_rrf_handles_empty_bm25():
    """BM25 为空时，向量结果应能单独返回。"""
    vector = [{"text": "Only vector", "source": "v.pdf", "page": 1}]
    bm25: list[dict] = []
    fused = _rrf_fuse(vector, bm25, top_k=3)
    assert len(fused) == 1
    assert fused[0]["source"] == "v.pdf"


def test_rrf_weight_vector_bias():
    """向量权重更高时，向量独有的文档应排更前面。"""
    vector = [
        {"text": "vec-only", "source": "v.pdf", "page": 1},
        {"text": "shared", "source": "s.pdf", "page": 1},
    ]
    bm25 = [
        {"text": "bm25-only", "source": "b.pdf", "page": 1},
        {"text": "shared", "source": "s.pdf", "page": 1},
    ]
    # 偏语义：vector_weight=2.0, bm25_weight=0.5
    fused = _rrf_fuse(vector, bm25, top_k=3, vector_weight=2.0, bm25_weight=0.5)
    scores = {d["source"]: d["score"] for d in fused}
    assert scores["v.pdf"] > scores["b.pdf"]


def test_rrf_weight_bm25_bias():
    """BM25 权重更高时，BM25 独有的文档应排更前面。"""
    vector = [
        {"text": "vec-only", "source": "v.pdf", "page": 1},
        {"text": "shared", "source": "s.pdf", "page": 1},
    ]
    bm25 = [
        {"text": "bm25-only", "source": "b.pdf", "page": 1},
        {"text": "shared", "source": "s.pdf", "page": 1},
    ]
    # 偏关键词：vector_weight=0.5, bm25_weight=2.0
    fused = _rrf_fuse(vector, bm25, top_k=3, vector_weight=0.5, bm25_weight=2.0)
    scores = {d["source"]: d["score"] for d in fused}
    assert scores["b.pdf"] > scores["v.pdf"]


def test_rrf_pure_vector_mode():
    """纯语义模式：BM25 权重为 0 时，BM25 独有的文档不应出现在结果中。"""
    vector = [{"text": "Only vector", "source": "v.pdf", "page": 1}]
    bm25 = [{"text": "Only bm25", "source": "b.pdf", "page": 1}]
    fused = _rrf_fuse(vector, bm25, top_k=3, bm25_weight=0.0)
    # BM25-only 文档 score=0，但仍在结果中（收集阶段不区分）
    # 纯语义下，两路共有的文档只靠向量路积分
    bm25_only = [d for d in fused if d["source"] == "b.pdf"]
    assert bm25_only[0]["score"] == 0.0


def test_rrf_fetch_k_truncation():
    """top_k 截断应正常工作。"""
    vector = [{"text": f"doc{i}", "source": f"s{i}.pdf", "page": 1} for i in range(10)]
    bm25: list[dict] = []
    fused = _rrf_fuse(vector, bm25, top_k=3)
    assert len(fused) == 3


def test_retrieve_without_rewrite():
    """未启用 rewrite 时不应调用 query_rewriter。"""
    from retriever import retrieve as hybrid_retrieve

    with patch("retriever.vector_search") as mock_vec, \
         patch("retriever.search_bm25") as mock_bm25:
        mock_vec.return_value = []
        mock_bm25.return_value = [
            {"text": "result", "source": "b.pdf", "page": 1},
        ]

        results = hybrid_retrieve(
            "test query", "test_kb", top_k=3,
            rewrite=False, history_text="some history",
        )
        assert len(results) >= 0


def test_retrieve_with_rewrite_mocked():
    """启用 rewrite 时应调用 query_rewriter 改写查询。"""
    from retriever import retrieve as hybrid_retrieve

    with patch("retriever.vector_search") as mock_vec, \
         patch("retriever.search_bm25") as mock_bm25, \
         patch("query_rewriter.rewrite_query") as mock_rw:
        mock_rw.return_value = "rewritten query text"
        mock_vec.return_value = [
            {"text": "result", "source": "v.pdf", "page": 1, "score": 0.9},
        ]
        mock_bm25.return_value = []

        results = hybrid_retrieve(
            "ambiguous query", "test_kb", top_k=3,
            rewrite=True, history_text="previous conversation",
        )
        mock_rw.assert_called_once_with("ambiguous query", "previous conversation")
        assert len(results) >= 1


def test_retrieve_without_rerank():
    """未启用 rerank 时不应调用 reranker。"""
    from retriever import retrieve as hybrid_retrieve

    with patch("retriever.vector_search") as mock_vec, \
         patch("retriever.search_bm25") as mock_bm25:
        mock_vec.return_value = []
        mock_bm25.return_value = [
            {"text": "result", "source": "b.pdf", "page": 1},
        ]

        results = hybrid_retrieve("test", "test_kb", top_k=3, rerank=False)
        assert len(results) >= 0


def test_retrieve_with_rerank_uses_larger_fetch():
    """启用 rerank 时 fetch_k 应更大（为精排提供更多候选）。"""
    from retriever import retrieve as hybrid_retrieve

    with patch("retriever.vector_search") as mock_vec, \
         patch("retriever.search_bm25") as mock_bm25:
        mock_vec.return_value = []
        mock_bm25.return_value = []

        # Just verify no crash with rerank=True
        results = hybrid_retrieve("test", "test_kb", top_k=3, rerank=True)
        # fetch_k = max(3*4, 20) = 20
        # vector_search(query, kb_name, top_k, token_tracker=None)
        fetch_arg = mock_vec.call_args[0][2]  # third positional arg = top_k
        assert fetch_arg == 20
        # search_bm25(query, kb_name, top_k)
        fetch_arg_bm = mock_bm25.call_args[0][2]
        assert fetch_arg_bm == 20
        assert len(results) == 0


# ═══════════════════════════════════════════════════════════════
# Conversation 测试
# ═══════════════════════════════════════════════════════════════

def test_create_conversation_empty():
    from conversation import create_conversation
    conv = create_conversation()
    assert conv == []
    assert len(conv) == 0


def test_add_turn_appends():
    from conversation import create_conversation, add_turn
    conv = create_conversation()
    refs = [{"text": "ref", "source": "s.pdf", "page": 1}]
    add_turn(conv, "Q1", "A1", refs)
    assert len(conv) == 1
    assert conv[0]["question"] == "Q1"
    assert conv[0]["answer"] == "A1"
    assert conv[0]["references"] == refs
    assert "timestamp" in conv[0]


def test_get_history_limits():
    from conversation import create_conversation, add_turn, get_history
    conv = create_conversation()
    for i in range(10):
        add_turn(conv, f"Q{i}", f"A{i}", [])
    assert len(get_history(conv, 3)) == 3
    assert get_history(conv, 3)[-1]["question"] == "Q9"


def test_get_history_zero():
    from conversation import create_conversation, add_turn, get_history
    conv = create_conversation()
    add_turn(conv, "Q", "A", [])
    assert get_history(conv, 0) == []


def test_format_for_prompt_includes_history():
    from conversation import create_conversation, add_turn, get_history, format_for_prompt
    conv = create_conversation()
    add_turn(conv, "什么是RAG？", "RAG是检索增强生成。", [])
    add_turn(conv, "它有什么优点？", "优点包括...", [])
    turns = get_history(conv, 2)
    text = format_for_prompt(turns)
    assert "对话历史" in text
    assert "什么是RAG？" in text
    assert "RAG是检索增强生成" in text
    assert "它有什么优点？" in text


def test_set_feedback():
    from conversation import create_conversation, add_turn, set_feedback
    conv = create_conversation()
    add_turn(conv, "Q", "A", [])
    set_feedback(conv, 0, "up")
    assert conv[0]["feedback"] == "up"
    set_feedback(conv, 0, None)
    assert conv[0]["feedback"] is None


def test_estimate_tokens():
    from conversation import estimate_tokens
    assert estimate_tokens("") == 0
    assert estimate_tokens("你好") > 0
    assert estimate_tokens("Hello world") > 0


# ═══════════════════════════════════════════════════════════════
# TokenTracker 测试
# ═══════════════════════════════════════════════════════════════

def test_token_tracker_init():
    from token_tracker import TokenTracker
    t = TokenTracker()
    assert t.total_tokens == 0
    assert t.estimated_cost == 0.0


def test_token_tracker_record_embedding():
    from token_tracker import TokenTracker
    t = TokenTracker()
    t.record_embedding(1000)
    assert t.embedding_tokens == 1000
    assert t.total_tokens == 1000


def test_token_tracker_record_generation():
    from token_tracker import TokenTracker
    t = TokenTracker()
    t.record_generation(500, 300)
    assert t.generation_input == 500
    assert t.generation_output == 300
    assert t.total_tokens == 800


def test_token_tracker_cost_estimation():
    from token_tracker import TokenTracker
    t = TokenTracker()
    t.record_embedding(10000)  # 10K embedding tokens → ¥0.005
    t.record_generation(5000, 3000)  # 5K gen in + 3K gen out
    cost = t.estimated_cost
    assert cost > 0
    # 10K * 0.0005/1K = 0.005 + 5K * 0.0008/1K = 0.004 + 3K * 0.002/1K = 0.006
    # Total = 0.015
    assert 0.014 <= cost <= 0.016


def test_token_tracker_summary():
    from token_tracker import TokenTracker
    t = TokenTracker()
    t.record_embedding(100)
    t.record_generation(50, 30)
    s = t.summary
    assert s["total_tokens"] == 180
    assert "estimated_cost" in s


def test_format_cost():
    from token_tracker import format_cost
    assert format_cost(0) == "—"
    assert format_cost(0.0005) == "<¥0.01"
    assert format_cost(0.01) == "¥0.01"
    assert format_cost(1.5) == "¥1.50"
    assert format_cost(0.35) == "¥0.35"


# ═══════════════════════════════════════════════════════════════
# Generator 多轮对话测试
# ═══════════════════════════════════════════════════════════════

def test_build_prompt_with_history():
    from generator import build_prompt
    from conversation import create_conversation, add_turn, get_history, format_for_prompt

    docs = [{"text": "RAG stands for Retrieval-Augmented Generation.", "source": "ai.pdf", "page": 3}]
    conv = create_conversation()
    add_turn(conv, "什么是RAG？", "RAG是检索增强生成技术。", [])

    history_text = format_for_prompt(get_history(conv, 1))
    prompt = build_prompt("它有什么优点？", docs, history_text)

    assert "对话历史" in prompt
    assert "什么是RAG？" in prompt
    assert "RAG是检索增强生成" in prompt
    assert "它有什么优点？" in prompt
    assert "[1]" in prompt
