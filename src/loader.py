from __future__ import annotations

import fitz
import requests
from bs4 import BeautifulSoup

from config import URL_TIMEOUT, USER_AGENT

SUPPORTED_TYPES = {
    ".pdf": "PDF 文档",
    ".txt": "纯文本",
    ".md": "Markdown",
    ".docx": "Word 文档",
}


# ── 表格转文本 ──────────────────────────────────────────────────────
def _table_to_markdown(table_data: list[list[str]]) -> str:
    """将二维表格转为 Markdown 风格文本，便于 LLM 理解和关键词匹配。"""
    if not table_data or not any(row for row in table_data):
        return ""

    # 过滤全空行
    rows = [[str(c or "") for c in row] for row in table_data if any(c for c in row)]
    if not rows:
        return ""

    col_count = max(len(row) for row in rows)
    # 补齐列数
    for row in rows:
        while len(row) < col_count:
            row.append("")

    lines = ["[表格]"]
    lines.append(" | ".join(rows[0]))

    if len(rows) > 1:
        lines.append(" | ".join("---" for _ in range(col_count)))
        for row in rows[1:]:
            lines.append(" | ".join(row))

    lines.append("[/表格]")
    return "\n".join(lines)


# ── PDF ─────────────────────────────────────────────────────────────
def _detect_scanned_page(page: fitz.Page, text: str) -> bool:
    """检测页面是否为扫描件/图片型：文字极少但有嵌入图片。"""
    text_len = len(text.strip())
    if text_len > 80:
        return False
    # 检查页面中是否有图片
    images = page.get_images(full=True)
    return len(images) > 0 and text_len < 80


def load_pdf(file_path: str) -> list[dict]:
    """提取 PDF 文本和表格，按页返回带元数据的列表。

    每页返回:
      {"page": int, "text": str, "tables": list, "is_scanned": bool}
    """
    doc = fitz.open(file_path)
    pages: list[dict] = []

    for i, page in enumerate(doc):
        page_num = i + 1

        # 表格提取
        tables_raw: list[list[list[str]]] = []
        try:
            found = page.find_tables()
            if found and found.tables:
                for t in found.tables:
                    extracted = t.extract()
                    if extracted:
                        tables_raw.append(extracted)
        except Exception:
            pass

        # 文本提取
        raw_text = page.get_text().strip()

        # 扫描件检测
        is_scanned = _detect_scanned_page(page, raw_text)

        # 组装页面文本：正文 + 表格 Markdown + 扫描警告
        text_parts: list[str] = []
        if raw_text:
            text_parts.append(raw_text)
        if is_scanned:
            text_parts.append("[图片型页面，文字内容可能不完整，建议使用 OCR 预处理]")

        for table_data in tables_raw:
            md = _table_to_markdown(table_data)
            if md:
                text_parts.append(md)

        combined_text = "\n\n".join(text_parts)

        if combined_text.strip() or tables_raw:
            pages.append({
                "page": page_num,
                "text": combined_text,
                "tables": tables_raw,
                "table_count": len(tables_raw),
                "is_scanned": is_scanned,
            })
        # 完全空页（无文字、无表格）跳过

    doc.close()
    return pages


# ── TXT / MD ────────────────────────────────────────────────────────
def load_txt(file_path: str) -> list[dict]:
    """读取纯文本文件，返回单页结构。

    优先 UTF-8，失败回退到 GBK（兼容中文 Windows 文档）。
    """
    for encoding in ("utf-8", "gbk"):
        try:
            with open(file_path, "r", encoding=encoding) as f:
                text = f.read().strip()
            break
        except (UnicodeDecodeError, UnicodeError):
            continue
    else:
        raise RuntimeError(f"无法解码文件: {file_path}")

    if not text:
        return []
    return [{"page": 1, "text": text, "tables": [], "table_count": 0, "is_scanned": False}]


# ── DOCX ────────────────────────────────────────────────────────────
def load_docx(file_path: str) -> list[dict]:
    """读取 Word (.docx) 文档，提取段落文本和表格。"""
    from docx import Document

    doc = Document(file_path)

    # 段落文本
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # 表格提取
    tables_raw: list[list[list[str]]] = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        # 过滤全空行
        non_empty = [r for r in rows if any(c for c in r)]
        if non_empty:
            tables_raw.append(non_empty)

    text_parts: list[str] = []
    if paragraphs:
        text_parts.append("\n\n".join(paragraphs))
    for t in tables_raw:
        md = _table_to_markdown(t)
        if md:
            text_parts.append(md)

    combined = "\n\n".join(text_parts)
    if not combined.strip():
        return []

    return [{
        "page": 1,
        "text": combined,
        "tables": tables_raw,
        "table_count": len(tables_raw),
        "is_scanned": False,
    }]


# ── URL ─────────────────────────────────────────────────────────────
def load_url(url: str) -> list[dict]:
    """抓取网页内容，提取正文文本。"""
    headers = {"User-Agent": USER_AGENT}
    resp = requests.get(url, headers=headers, timeout=URL_TIMEOUT)
    resp.raise_for_status()
    resp.encoding = resp.apparent_encoding

    soup = BeautifulSoup(resp.text, "html.parser")

    for tag in soup(["script", "style", "nav", "footer", "header", "noscript"]):
        tag.decompose()

    title = soup.title.string.strip() if soup.title and soup.title.string else url
    body = soup.body.get_text(separator="\n", strip=True) if soup.body else soup.get_text(separator="\n", strip=True)

    text = f"标题: {title}\n\n{body}"

    if not text.strip():
        return []
    return [{"page": 1, "text": text, "tables": [], "table_count": 0, "is_scanned": False}]


# ── 调度器 ─────────────────────────────────────────────────────────
def load_document(file_path: str, doc_type: str) -> list[dict]:
    """根据文档类型调度对应的 loader。"""
    if doc_type == "url":
        return load_url(file_path)
    if doc_type == ".pdf":
        return load_pdf(file_path)
    if doc_type == ".txt" or doc_type == ".md":
        return load_txt(file_path)
    if doc_type == ".docx":
        return load_docx(file_path)
    raise ValueError(f"不支持的文档类型: {doc_type}")
